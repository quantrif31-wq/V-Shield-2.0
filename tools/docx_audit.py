import json, sys
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield 'P', Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield 'T', Table(child, doc)

def main():
    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    doc = Document(src)
    rows=[]
    for i,(kind,obj) in enumerate(blocks(doc)):
        if kind=='P':
            rows.append({'i':i,'kind':'P','style':obj.style.name if obj.style else '', 'text':obj.text})
        else:
            rows.append({'i':i,'kind':'T','rows':[[c.text for c in r.cells] for r in obj.rows]})
    out.write_text(json.dumps({'paragraphs':len(doc.paragraphs),'tables':len(doc.tables),'sections':len(doc.sections),'blocks':rows},ensure_ascii=False,indent=2),encoding='utf-8')

if __name__=='__main__': main()
