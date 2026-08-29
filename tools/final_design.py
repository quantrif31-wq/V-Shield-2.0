from pathlib import Path
from PIL import Image,ImageDraw,ImageFont
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE=Path(r'C:\Users\AnhPhamPC\Downloads\PRO2191 - Dự án TN (UDPM-.NET Core)')
SRC=BASE/'BaoCaoTotNghiep_VShield_Chapter3_4_Updated.docx'
OUT=BASE/'BaoCaoTotNghiep_VShield_Chapter3_4_FinalDesign.docx'
AS=BASE/'vshield_diagrams_final'; AS.mkdir(exist_ok=True)
W,H=2400,1500; F=r'C:\Windows\Fonts\arial.ttf'; FB=r'C:\Windows\Fonts\arialbd.ttf'
C={'ink':'#173247','blue':'#245E82','line':'#6B8799','pale':'#EAF3F8','pale2':'#F6F9FB','red':'#B43A40','redp':'#FBEAEC','green':'#287A55','greenp':'#E8F5EE','amber':'#B87418','amberp':'#FFF4DD','white':'#FFFFFF','gray':'#596B78'}
def ft(n,b=False): return ImageFont.truetype(FB if b else F,n)
def wrap(d,s,f,mw):
    out=[]; cur=''
    for w in str(s).split():
        t=(cur+' '+w).strip()
        if d.textbbox((0,0),t,font=f)[2]<=mw: cur=t
        else:
            if cur: out.append(cur)
            cur=w
    if cur: out.append(cur)
    return out
def txt(d,xy,s,size=28,b=False,color=None,maxw=None,anchor=None):
    f=ft(size,b); x,y=xy
    if maxw:
        for line in wrap(d,s,f,maxw): d.text((x,y),line,font=f,fill=color or C['ink'],anchor=anchor); y+=size*1.25
    else:d.text((x,y),s,font=f,fill=color or C['ink'],anchor=anchor)
def rect(d,xy,fill='white',outline=None,r=18,w=3): d.rounded_rectangle(xy,r,fill=fill,outline=outline or C['line'],width=w)
def line(d,a,b,color=None,w=4,dash=False,arrow=False):
    color=color or C['line']
    if dash:
        import math
        x1,y1=a;x2,y2=b;L=math.hypot(x2-x1,y2-y1); ux=(x2-x1)/L;uy=(y2-y1)/L
        z=0
        while z<L:
            z2=min(z+18,L);d.line((x1+ux*z,y1+uy*z,x1+ux*z2,y1+uy*z2),fill=color,width=w);z+=30
    else:d.line((a,b),fill=color,width=w)
    if arrow:
        x,y=b; d.polygon([(x,y),(x-18,y-10),(x-18,y+10)],fill=color)
def base(title,sub=None):
    im=Image.new('RGB',(W,H),'white');d=ImageDraw.Draw(im);txt(d,(70,45),title,44,True,C['ink']);d.line((70,108,W-70,108),fill=C['blue'],width=6)
    if sub:txt(d,(W-70,60),sub,23,False,C['gray'],anchor='ra')
    return im,d
def save(im,name):
    p=AS/name;im.save(p,dpi=(240,240),quality=96);return p
def actor(d,x,y,label):
    d.ellipse((x-22,y-75,x+22,y-31),outline=C['ink'],width=5);d.line((x,y-31,x,y+30),fill=C['ink'],width=5);d.line((x-42,y-5,x+42,y-5),fill=C['ink'],width=5);d.line((x,y+30,x-35,y+78),fill=C['ink'],width=5);d.line((x,y+30,x+35,y+78),fill=C['ink'],width=5);txt(d,(x,y+90),label,23,True,anchor='ma')
def ellipse(d,c,label,w=340,h=86,fill=None):
    x,y=c;d.ellipse((x-w/2,y-h/2,x+w/2,y+h/2),fill=fill or C['pale2'],outline=C['blue'],width=4);txt(d,(x,y),label,23,False,anchor='mm',maxw=None)
def boundary(d,xy,title):
    x1,y1,x2,y2=xy;d.rectangle(xy,outline=C['blue'],width=5);rect(d,(x1+20,y1+15,x1+500,y1+70),C['blue'],C['blue'],10);txt(d,(x1+40,y1+28),title,26,True,'white')
def rel(d,a,b,label=None,dash=False):
    line(d,a,b,dash=dash,arrow=dash)
    if label:txt(d,((a[0]+b[0])//2,(a[1]+b[1])//2-28),label,21,True,C['blue'],anchor='mm')
def panel(d,xy,title,items=(),accent=None):
    x1,y1,x2,y2=xy;rect(d,xy,C['white'],accent or C['line'],16,3);d.rectangle((x1,y1,x2,y1+58),fill=accent or C['pale']);txt(d,(x1+20,y1+14),title,25,True,'white' if accent else C['ink'])
    y=y1+78
    for it in items:txt(d,(x1+22,y),it,22,False,C['ink'],x2-x1-44);y+=42
def badge(d,xy,label,color,pale):rect(d,xy,pale,color,18,2);txt(d,((xy[0]+xy[2])//2,(xy[1]+xy[3])//2),label,21,True,color,anchor='mm')
def button(d,xy,label,kind='primary'):
    colors={'primary':(C['blue'],C['blue'],'white'),'danger':(C['redp'],C['red'],C['red']),'success':(C['green'],C['green'],'white'),'neutral':('white',C['line'],C['ink'])};f,o,t=colors[kind];rect(d,xy,f,o,12,3);txt(d,((xy[0]+xy[2])//2,(xy[1]+xy[3])//2),label,23,True,t,anchor='mm')
def call(d,x,y,n):d.ellipse((x-25,y-25,x+25,y+25),fill=C['blue'],outline='white',width=3);txt(d,(x,y),str(n),23,True,'white',anchor='mm')

def system_context():
 im,d=base('Mô hình tổng quan hệ thống V-Shield 2.0','System context / high-level architecture')
 # users
 panel(d,(65,155,380,1365),'NGƯỜI DÙNG',['Admin','Security Manager','Bảo vệ cổng','Bảo vệ cơ động','Lễ tân','Nhân viên','Visitor'],C['blue'])
 for i in range(7): actor(d,135,300+i*140,''); line(d,(180,225+i*140),(430,300+i*140),w=2)
 panel(d,(430,155,860,1365),'CLIENT CHANNELS',['Web Admin / SOC','Gate Transit Console','Android Mobile App','Visitor Kiosk'],C['blue'])
 for i,lab in enumerate(['WEB','GATE','MOBILE','KIOSK']):rect(d,(485,310+i*230,805,460+i*230),C['pale'],C['blue'],20);txt(d,(645,385+i*230),lab,28,True,C['blue'],anchor='mm')
 boundary(d,(930,155,1840,1365),'V-SHIELD 2.0')
 modules=['REST API','Authentication / RBAC','Gate Transit','SOC','Visitor','HR / Attendance','Vehicle','Evidence','Notification','SignalR Hub']
 for i,m in enumerate(modules):
  x=990+(i%2)*410;y=270+(i//2)*190;rect(d,(x,y,x+355,y+112),C['pale2'],C['blue'],18);txt(d,(x+177,y+56),m,23,True,anchor='mm')
 panel(d,(1910,155,2335,1365),'INFRA / EDGE',['SQL Server','AI / ANPR Edge','Camera','QR Scanner','Barrier'],C['gray'])
 for i in range(5):rect(d,(1975,330+i*175,2270,425+i*175),C['pale2'],C['line'],12)
 line(d,(860,760),(930,760),C['blue'],6,arrow=True);line(d,(1840,760),(1910,760),C['blue'],6,arrow=True)
 return save(im,'fig_3_1_system_context.png')
def usecase(name,title,btitle,actors,uses,links,rels=()):
 im,d=base(title,'UML Use Case Diagram');boundary(d,(480,165,1930,1360),btitle)
 positions={}
 for i,u in enumerate(uses):
  x=810+(i%3)*440;y=310+(i//3)*185;positions[u]=(x,y);ellipse(d,(x,y),u,360,88)
 apos={}
 left=[a for a,s in actors if s=='L'];right=[a for a,s in actors if s=='R']
 for i,a in enumerate(left):apos[a]=(210,300+i*(950//max(1,len(left))));actor(d,*apos[a],a)
 for i,a in enumerate(right):apos[a]=(2180,300+i*(950//max(1,len(right))));actor(d,*apos[a],a)
 for a,u in links:
  ax,ay=apos[a];ux,uy=positions[u];rel(d,(ax+70 if ax<480 else ax-70,ay),(ux-180 if ax<480 else ux+180,uy))
 for src,dst,label in rels:
  a=positions[src];b=positions[dst];rel(d,a,b,label,True)
 return save(im,name)
def architecture():
 im,d=base('Kiến trúc phân lớp V-Shield 2.0','ASP.NET Core 8 Modular Monolith')
 layers=[('1. PRESENTATION',['Web Admin / SOC\nVue 3','Gate Transit Console\nVue 3','Android App\nKotlin','Visitor Kiosk']),('2. API / COMMUNICATION',['HTTPS REST\nrequest / response','JWT / RBAC\nauthentication','SignalR\nrealtime events']),('3. APPLICATION',['Auth','Gate Transit','SOC','Visitor','HR / Attendance','Vehicle','Evidence','Notification']),('4. DATA / INTEGRATION',['Entity Framework Core','SQL Server','Edge AI / ANPR','Camera','QR Reader','Barrier'])]
 y=155
 for li,(head,items) in enumerate(layers):
  h=245 if li!=2 else 300;rect(d,(75,y,2325,y+h),C['pale2'] if li%2==0 else 'white',C['blue'],16);rect(d,(95,y+25,390,y+h-25),C['blue'],C['blue'],12);txt(d,(242,y+h/2),head,25,True,'white',anchor='mm')
  gap=28;cw=(1880-gap*(len(items)-1))//len(items)
  for i,it in enumerate(items):x=420+i*(cw+gap);rect(d,(x,y+55,x+cw,y+h-55),C['white'],C['line'],14);txt(d,(x+cw/2,y+h/2),it,23,True,anchor='mm')
  if li<3:line(d,(1200,y+h),(1200,y+h+45),C['blue'],6,arrow=True)
  y+=h+55
 return save(im,'fig_4_1_architecture.png')
def sitemap():
 im,d=base('Sitemap V-Shield 2.0','Tree navigation model');rect(d,(980,145,1420,225),C['blue'],C['blue'],18);txt(d,(1200,185),'V-SHIELD',30,True,'white',anchor='mm')
 roots=[('WEB ADMIN / SOC',['Dashboard','SOC','Gate Control','Visitors','HR & Attendance','Evidence','Devices','Administration']),('MOBILE',['Home','My QR','Alerts','Incident','Visitor','Vehicle','Attendance','Leave','Profile']),('KIOSK',['Welcome','Scan QR','Check-in','Confirmation'])]
 xs=[430,1200,1970]
 details={'SOC':'Active Alarms · Incident Detail · SOP','Gate Control':'Gate Transit · Barrier · History','Visitors':'Pre-registration · Pass · Check-in · Overstay · Watchlist','HR & Attendance':'Employees · Shift · Attendance · Anomaly · Leave','Evidence':'Repository · Verification · Redaction · Export · Custody'}
 for (root,items),x in zip(roots,xs):
  line(d,(1200,225),(x,300),C['line'],4);rect(d,(x-250,270,x+250,340),C['pale'],C['blue'],14);txt(d,(x,305),root,25,True,C['blue'],anchor='mm')
  y=390
  for it in items:
   line(d,(x-220,y+28),(x-150,y+28),C['line'],3);rect(d,(x-150,y,x+250,y+56),'white',C['line'],10);txt(d,(x-130,y+15),it,21,True)
   if it in details:txt(d,(x-130,y+62),details[it],17,False,C['gray'],400);y+=55
   y+=82
 return save(im,'fig_4_2_sitemap.png')
def desktop_shell(d,title='Dashboard'):
 rect(d,(55,140,2345,1390),'white',C['line'],18);d.rectangle((55,140,2345,245),fill=C['ink']);txt(d,(90,170),'V-SHIELD',30,True,'white');badge(d,(430,167,640,220),'SITE HCM',C['blue'],C['pale']);badge(d,(1730,167,1935,220),'SYSTEM ONLINE',C['green'],C['greenp']);txt(d,(2060,182),'🔔   Admin',24,True,'white')
 d.rectangle((55,245,390,1390),fill='#F1F5F8');
 for i,m in enumerate(['Dashboard','SOC','Gate','Visitor','HR','Evidence','Devices','Administration']):
  if m==title:rect(d,(78,300+i*92,368,365+i*92),C['pale'],C['blue'],10)
  txt(d,(105,318+i*92),m,23,True if m==title else False,C['blue'] if m==title else C['ink'])
 txt(d,(435,285),title,34,True)
def layout_admin():
 im,d=base('Layout 1 - Web Admin / SOC','Desktop 16:9 low-fidelity wireframe');desktop_shell(d,'SOC')
 for i,(lab,val) in enumerate([('CRITICAL','03'),('ACTIVE','18'),('OFFLINE','02')]):panel(d,(435+i*330,350,735+i*330,510),lab,[val],C['red'] if i==0 else C['blue'])
 rect(d,(1450,350,2260,510),'white',C['line'],12);txt(d,(1480,375),'FILTER',20,True);txt(d,(1480,425),'Severity  ▾    Status  ▾    Today  ▾',22)
 panel(d,(435,550,1580,1315),'MAIN WORKSPACE',['Alarm / incident list','Table, pagination, selection'],C['blue']);panel(d,(1615,550,2260,1315),'ACTIVITY',['Realtime events','Assigned tasks','Audit timeline'],C['gray']);return save(im,'fig_4_3_admin_wireframe.png')
def gate_ui(name,title,calls=False):
 im,d=base(title,'Gate Transit operational console');rect(d,(60,145,2340,1390),'white',C['line'],18);d.rectangle((60,145,2340,245),fill=C['ink']);txt(d,(95,175),'V-SHIELD / GATE 01',29,True,'white');badge(d,(1660,170,1840,220),'ONLINE',C['green'],C['greenp']);txt(d,(2060,180),'15:31:08',25,True,'white')
 panel(d,(95,280,1450,995),'LIVE CAMERA',['CAM-ENTRY-01   •   1920×1080'],C['blue']);d.rectangle((140,390,1405,870),fill='#DCE5EA');txt(d,(772,630),'CAMERA 16:9',48,True,C['gray'],anchor='mm');badge(d,(160,900,460,960),'51A-123.45',C['blue'],C['pale']);txt(d,(500,915),'ANPR confidence 98.7%',24,True)
 panel(d,(1490,280,2305,720),'IDENTITY',['Nguyễn Văn A','Employee · IT Department','QR: VALID'],C['blue']);badge(d,(1990,390,2250,445),'QR VALID',C['green'],C['greenp']);panel(d,(1490,750,2305,995),'VEHICLE',['51A-123.45 · Toyota','Owner: Nguyễn Văn A'],C['gray'])
 rect(d,(95,1030,1450,1335),C['greenp'],C['green'],18,5);txt(d,(772,1115),'ALLOW',58,True,C['green'],anchor='mm');txt(d,(772,1195),'Identity, QR and vehicle matched',25,False,C['green'],anchor='mm');button(d,(1530,1045,2260,1135),'OPEN BARRIER','success');button(d,(1530,1160,1955,1250),'OVERRIDE','neutral');button(d,(1990,1160,2260,1250),'DURESS','danger')
 if calls:
  for x,y,n in [(120,310,1),(180,920,2),(2190,420,3),(1520,310,4),(1380,1080,5),(2220,1090,6),(1915,1200,7),(2230,1200,8)]:call(d,x,y,n)
 return save(im,name)
def phone(name,title,calls=False):
 im,d=base(title,'Android portrait wireframe');x1,x2=780,1620;rect(d,(x1,145,x2,1400),'#F7F9FA',C['ink'],55,8);d.rectangle((x1+35,190,x2-35,300),fill=C['ink']);txt(d,(x1+70,220),'V-Shield',28,True,'white');txt(d,(x2-120,220),'●  🔔',25,False,'white')
 rect(d,(x1+70,350,x2-70,740),C['redp'],C['red'],26,5);badge(d,(x1+110,390,x1+360,450),'CRITICAL',C['red'],C['redp']);txt(d,(x1+110,490),'Duress Alarm',38,True,C['red']);txt(d,(x1+110,555),'Gate 02 · 15:31',28,True);txt(d,(x1+110,610),'Source: Gate Console\nStatus: New',24,False,C['gray'])
 panel(d,(x1+70,780,x2-70,1035),'INCIDENT DETAIL',['Bảo vệ cổng kích hoạt cảnh báo cưỡng ép.','Ưu tiên phản hồi ngay.'],C['blue']);button(d,(x1+90,1080,x2-90,1170),'ACKNOWLEDGE','primary');button(d,(x1+90,1190,x2-90,1280),'OPEN INCIDENT','neutral');txt(d,((x1+x2)//2,1340),'Home      Alerts      Chat      Me',22,True,C['gray'],anchor='mm')
 if calls:
  for x,y,n in [(830,390,1),(830,510,2),(830,590,3),(830,810,4),(1580,1125,5),(1580,1235,6)]:call(d,x,y,n)
 return save(im,name)
def kiosk():
 im,d=base('Layout 4 - Visitor Kiosk','Portrait / tablet wireframe');rect(d,(650,145,1750,1390),'#F7FAFC',C['ink'],35,6);txt(d,(1200,240),'V-SHIELD',36,True,C['blue'],anchor='mm');txt(d,(1200,325),'WELCOME',31,True,anchor='mm');txt(d,(1200,380),'Quét QR để check-in',25,False,C['gray'],anchor='mm');rect(d,(835,470,1565,970),'white',C['blue'],24,5);d.line((900,535,1080,535),fill=C['blue'],width=8);d.line((900,535,900,715),fill=C['blue'],width=8);d.line((1500,535,1320,535),fill=C['blue'],width=8);d.line((1500,535,1500,715),fill=C['blue'],width=8);txt(d,(1200,735),'QR SCANNER',35,True,C['gray'],anchor='mm');button(d,(835,1025,1565,1120),'SCAN QR','primary');button(d,(835,1150,1565,1245),'MANUAL CHECK-IN','neutral');txt(d,(1200,1320),'VI  |  EN     •     Cần trợ giúp?',21,False,C['gray'],anchor='mm');return save(im,'fig_4_6_kiosk_wireframe.png')
def soc():
 im,d=base('Screen 01 - SOC Alarm Console','Functional wireframe with callouts');desktop_shell(d,'SOC')
 for i,(lab,val,col) in enumerate([('CRITICAL','03','red'),('HIGH','07','amber'),('ACTIVE','18','blue'),('DEVICE OFFLINE','02','gray')]):x=430+i*445;panel(d,(x,335,x+410,505),lab,[val],C[col]);call(d,x+25,360,1)
 panel(d,(430,540,1020,1315),'ACTIVE ALARMS',['CRITICAL · 15:31 · Gate 02','HIGH · 15:22 · Camera 08','MEDIUM · 15:10 · Visitor'],C['red']);call(d,455,575,2)
 panel(d,(1050,540,2260,1010),'INCIDENT DETAIL',['Duress Alarm · Gate 02 · 15:31','Camera snapshot / evidence preview','Timeline: New → Acknowledged','Assigned: Bảo vệ cơ động 02'],C['blue']);call(d,1075,575,3);d.rectangle((1100,700,1580,930),fill='#DCE5EA');txt(d,(1340,815),'SNAPSHOT',30,True,C['gray'],anchor='mm')
 button(d,(1100,1080,1430,1170),'ACKNOWLEDGE','primary');button(d,(1470,1080,1770,1170),'ASSIGN','neutral');button(d,(1810,1080,2140,1170),'CLOSE','danger');call(d,1120,1085,4);call(d,1490,1085,5);call(d,1830,1085,6);return save(im,'fig_4_7_soc_alarm_console.png')
def visitor_screen():
 im,d=base('Screen 03 - Visitor Check-in','Kiosk workflow with callouts');steps=[('SCAN QR','QR-2026-0814'),('VALIDATION','QR valid · time window valid'),('VISITOR INFO','Trần Minh B · Host: Nguyễn Văn A'),('CHECK-IN','Expected 15:00–17:00'),('PASS GENERATED','VP-08421 · Gate 01')]
 for i,(h,s) in enumerate(steps):x=75+i*465;rect(d,(x,330,x+390,1010),C['white'],C['blue'],20);call(d,x+30,365,i+1);txt(d,(x+195,460),h,28,True,C['blue'],anchor='mm');d.rectangle((x+55,540,x+335,755),fill=C['pale']);txt(d,(x+195,650),'QR' if i==0 else '✓',48,True,C['green'],anchor='mm');txt(d,(x+35,810),s,22,False,C['ink'],320); 
 for i in range(4):line(d,(465+i*465,670),(535+i*465,670),C['blue'],5,arrow=True)
 button(d,(850,1120,1550,1220),'CONFIRM CHECK-IN','success');return save(im,'fig_4_9_visitor_kiosk.png')
def evidence():
 im,d=base('Screen 05 - Evidence Repository','Functional desktop wireframe');desktop_shell(d,'Evidence')
 for i,(a,b) in enumerate([('TOTAL EVIDENCE','1,284'),('PENDING VERIFY','12'),('EXPORT PENDING','05')]):panel(d,(430+i*430,340,830+i*430,500),a,[b],C['blue']);call(d,450+i*430,360,1)
 rect(d,(1745,340,2260,500),'white',C['line'],12);txt(d,(1770,365),'FILTER',20,True);txt(d,(1770,415),'Incident ▾  Type ▾  Integrity ▾',20)
 panel(d,(430,540,1740,1305),'EVIDENCE TABLE',['EV-2401 | INC-108 | Video | SHA-256 | Verified','EV-2402 | INC-109 | Image | SHA-256 | Pending','EV-2403 | INC-109 | Document | SHA-256 | Verified'],C['blue']);call(d,455,575,2)
 panel(d,(1770,540,2260,1305),'DETAIL DRAWER',['EV-2402','Hash: 4d8f...','Integrity: Pending','Custody: 3 events'],C['gray']);call(d,1795,575,3);button(d,(1810,900,2220,980),'VERIFY','primary');button(d,(1810,1000,2220,1080),'REDACTION','neutral');button(d,(1810,1100,2220,1180),'EXPORT','neutral');button(d,(1810,1200,2220,1280),'CUSTODY','neutral');call(d,2190,940,4);call(d,2190,1040,5);call(d,2190,1140,6);call(d,2190,1240,7);return save(im,'fig_4_11_evidence_repository.png')
def domain():
 im,d=base('Mô hình miền dữ liệu V-Shield 2.0','Domain dependencies — not a physical ERD');domains=[('IDENTITY',['AppUser','Employee']),('ACCESS & VEHICLE',['Gate','AccessLog','ZoneTransit','Vehicle','Delegation']),('VISITOR & HR',['PreRegistration','VisitorDetail','Attendance','LeaveRequest']),('SOC',['Alarm','Incident','DispatchTask']),('EVIDENCE',['EvidenceItem','Redaction','Export','Custody'])]
 pos=[(400,430),(1200,300),(1200,950),(2000,430),(2000,1050)]
 for (h,items),(x,y) in zip(domains,pos):panel(d,(x-300,y-170,x+300,y+170),h,items,C['blue'])
 for a,b in [(pos[0],pos[1]),(pos[0],pos[2]),(pos[1],pos[3]),(pos[3],pos[4]),(pos[1],pos[4])]:line(d,a,b,C['line'],5,arrow=True)
 return save(im,'fig_4_12_data_domain.png')
def entity(d,xy,name,attrs):
 x1,y1,x2,y2=xy;rect(d,xy,'white',C['blue'],10,3);d.rectangle((x1,y1,x2,y1+55),fill=C['blue']);txt(d,((x1+x2)//2,y1+27),name,22,True,'white',anchor='mm');y=y1+72
 for a in attrs:txt(d,(x1+15,y),a,18,True if a.startswith(('PK','FK')) else False,C['ink']);y+=32
def crow(d,a,b,label='1     N',logical=False):
 line(d,a,b,C['line'],3,dash=logical);x,y=b;d.line((x,y,x-18,y-15),fill=C['line'],width=3);d.line((x,y,x-18,y+15),fill=C['line'],width=3);txt(d,((a[0]+b[0])//2,(a[1]+b[1])//2-24),label,18,True,C['gray'],anchor='mm')
def erd(name,title,ents,rels):
 im,d=base(title,"Crow's-foot ERD · PK/FK from source");pos={}
 for key,xy,attrs in ents:entity(d,xy,key,attrs);pos[key]=xy
 for a,b,logical in rels:
  A=pos[a];B=pos[b];p1=(A[2],(A[1]+A[3])//2);p2=(B[0],(B[1]+B[3])//2);crow(d,p1,p2,'1        N',logical)
 if any(x[2] for x in rels):txt(d,(80,1410),'Dashed relationship = logical relationship; no SQL FK asserted.',20,False,C['gray'])
 return save(im,name)
def classdiag():
 im,d=base('UML Class Diagram - Gate Transit','Controller → Service → Data and entities')
 classes=[('GateTransitController',(70,230,700,730),['- zoneTransitService : ZoneTransitService'],['+ ScanAsync(...)','+ ScanGuestAsync(...)','+ OverrideAsync(...)','+ DuressAsync(...)']),('ZoneTransitService',(880,230,1510,800),['- dbContext : ApplicationDbContext'],['+ ProcessAccessLogAsync(id)','+ ProcessTransitAsync(...)','+ CreateTransitAsync(...)','+ GetTransitsAsync(...)','+ QueryTransitsAsync(...)']),('ApplicationDbContext',(1690,230,2330,800),['+ Gates : DbSet<Gate>','+ Vehicles : DbSet<Vehicle>','+ ZoneTransits : DbSet<ZoneTransit>','+ AccessLogs : DbSet<AccessLog>'],[])]
 for n,xy,attrs,methods in classes:
  x1,y1,x2,y2=xy;rect(d,xy,'white',C['blue'],8,4);d.rectangle((x1,y1,x2,y1+70),fill=C['blue']);txt(d,((x1+x2)//2,y1+35),n,25,True,'white',anchor='mm');mid=y1+70+max(120,len(attrs)*45);d.line((x1,mid,x2,mid),fill=C['blue'],width=3);y=y1+90
  for a in attrs:txt(d,(x1+20,y),a,20);y+=42
  y=mid+20
  for m in methods:txt(d,(x1+20,y),m,20);y+=42
 line(d,(700,480),(880,480),C['blue'],4,dash=True,arrow=True);txt(d,(790,440),'dependency',18,False,C['gray'],anchor='mm');line(d,(1510,480),(1690,480),C['blue'],4,dash=True,arrow=True)
 ents=[('Gate',(180,980,630,1270),['PK GateId','GateName','Location'],[]),('Vehicle',(710,980,1160,1310),['PK VehicleId','FK EmployeeId','LicensePlate'],[]),('ZoneTransit',(1240,930,1770,1340),['PK ZoneTransitId','FK EmployeeId','FK SecurityZoneId','FK AccessLogId?','Direction','Timestamp'],[]),('AccessLog',(1850,980,2300,1310),['PK AccessLogId','FK GateId','EmployeeId?'],[])]
 for n,xy,a,m in ents:
  entity(d,xy,n,a)
 line(d,(2010,800),(2070,980),C['line'],3);line(d,(2010,800),(1505,930),C['line'],3);line(d,(2010,800),(405,980),C['line'],3);line(d,(2010,800),(935,980),C['line'],3)
 return save(im,'fig_4_16_class_diagram.png')

assets={
'Hình 3.1.':system_context(),
'Hình 3.2.':usecase('fig_3_2_usecase_overview.png','Use Case tổng quan','V-SHIELD 2.0',[(a,'L') for a in ['Admin','QuanLy / NhanSu','BaoVe Gate','BaoVe Roving']]+[(a,'R') for a in ['LeTan','NhanVien','Visitor']],['Quản trị hệ thống','Kiểm soát vào/ra','Vận hành SOC','Quản lý khách','Quản lý phương tiện','HR & Attendance','Evidence & Compliance'],[('Admin','Quản trị hệ thống'),('Admin','Evidence & Compliance'),('QuanLy / NhanSu','HR & Attendance'),('QuanLy / NhanSu','Vận hành SOC'),('BaoVe Gate','Kiểm soát vào/ra'),('BaoVe Roving','Vận hành SOC'),('LeTan','Quản lý khách'),('NhanVien','Quản lý khách'),('NhanVien','Quản lý phương tiện'),('Visitor','Quản lý khách')]),
'Hình 3.3.':usecase('fig_3_3_usecase_gate_soc.png','Use Case Gate Transit & SOC','V-SHIELD – GATE TRANSIT & SOC',[('Bảo vệ cổng','L'),('Bảo vệ cơ động','L'),('Security Manager','R'),('ANPR Camera','R'),('Barrier','R')],['Kiểm soát lượt vào/ra','Scan QR','Nhận diện ANPR','Xác thực quyền','Ghi Transit / Receipt','Mở Barrier','Override mở cổng','Nhập lý do','Ghi Audit / Receipt','Kích hoạt Duress','Tạo Critical Alarm','Gửi realtime','Receive Alarm','Acknowledge','Assign Incident','Comment Incident','Close Incident'],[('Bảo vệ cổng','Kiểm soát lượt vào/ra'),('Bảo vệ cổng','Override mở cổng'),('Bảo vệ cổng','Kích hoạt Duress'),('Bảo vệ cơ động','Receive Alarm'),('Bảo vệ cơ động','Acknowledge'),('Bảo vệ cơ động','Comment Incident'),('Security Manager','Assign Incident'),('Security Manager','Close Incident'),('ANPR Camera','Nhận diện ANPR'),('Barrier','Mở Barrier')],[('Kiểm soát lượt vào/ra','Scan QR','<<include>>'),('Kiểm soát lượt vào/ra','Nhận diện ANPR','<<include>>'),('Kiểm soát lượt vào/ra','Xác thực quyền','<<include>>'),('Kiểm soát lượt vào/ra','Ghi Transit / Receipt','<<include>>'),('Override mở cổng','Kiểm soát lượt vào/ra','<<extend>>'),('Override mở cổng','Nhập lý do','<<include>>'),('Override mở cổng','Ghi Audit / Receipt','<<include>>'),('Kích hoạt Duress','Tạo Critical Alarm','<<include>>'),('Kích hoạt Duress','Gửi realtime','<<include>>')]),
'Hình 3.4.':usecase('fig_3_4_usecase_visitor.png','Use Case Visitor & Employee Self-Service','V-SHIELD – VISITOR & EMPLOYEE SERVICES',[('NhanVien','L'),('Visitor','L'),('LeTan','R'),('QuanLy / NhanSu','R')],['Pre-register Visitor','Create Token / Link','Self Registration','Enter Visitor Info','Kiosk Check-in','Validate QR','Validate Time Window','Approve Visitor','Issue Visitor Pass','Manual Check-in Support','Register Vehicle','Delegate Vehicle','Create Leave Request','View Attendance','Approve Leave Request'],[('NhanVien','Pre-register Visitor'),('NhanVien','Register Vehicle'),('NhanVien','Delegate Vehicle'),('NhanVien','Create Leave Request'),('NhanVien','View Attendance'),('Visitor','Self Registration'),('Visitor','Kiosk Check-in'),('LeTan','Approve Visitor'),('LeTan','Issue Visitor Pass'),('LeTan','Manual Check-in Support'),('QuanLy / NhanSu','Approve Leave Request')],[('Pre-register Visitor','Create Token / Link','<<include>>'),('Self Registration','Enter Visitor Info','<<include>>'),('Kiosk Check-in','Validate QR','<<include>>'),('Kiosk Check-in','Validate Time Window','<<include>>')]),
'Hình 3.5.':usecase('fig_3_5_usecase_admin.png','Use Case Administration, HR & Evidence','V-SHIELD – ENTERPRISE ADMINISTRATION',[('Admin','L'),('QuanLy','R'),('NhanSu','R')],['Manage User','Manage Role / RBAC','Device Topology','Notification Rules','Audit Logs','Full Lockdown','Manage Employee','Manage Shift','Calculate Attendance','Review Anomaly','Approve Leave','Register Evidence','Verify Integrity','Review Export','Redaction','Export Evidence','Chain of Custody'],[('Admin','Manage User'),('Admin','Manage Role / RBAC'),('Admin','Device Topology'),('Admin','Notification Rules'),('Admin','Audit Logs'),('Admin','Full Lockdown'),('Admin','Register Evidence'),('Admin','Verify Integrity'),('QuanLy','Review Export'),('QuanLy','Export Evidence'),('QuanLy','Chain of Custody'),('NhanSu','Manage Employee'),('NhanSu','Manage Shift'),('NhanSu','Calculate Attendance'),('NhanSu','Review Anomaly'),('NhanSu','Approve Leave')]),
'Hình 4.1.':architecture(),'Hình 4.2.':sitemap(),'Hình 4.3.':layout_admin(),'Hình 4.4.':gate_ui('fig_4_4_gate_wireframe.png','Layout 2 - Gate Transit Console'),'Hình 4.5.':phone('fig_4_5_mobile_wireframe.png','Layout 3 - Android Mobile App'),'Hình 4.6.':kiosk(),'Hình 4.7.':soc(),'Hình 4.8.':gate_ui('fig_4_8_gate_console.png','Screen 02 - Gate Transit Console',True),'Hình 4.9.':visitor_screen(),'Hình 4.10.':phone('fig_4_10_mobile_alert.png','Screen 04 - Mobile Critical Alert',True),'Hình 4.11.':evidence(),'Hình 4.12.':domain(),
'Hình 4.13.':erd('fig_4_13_erd_access_vehicle.png','ERD - Access & Vehicle',[('EMPLOYEE',(70,270,520,650),['PK EmployeeId','Name','Department']),('VEHICLE',(690,220,1160,670),['PK VehicleId','FK EmployeeId','LicensePlate','ParkingStatus']),('VEHICLE_DELEGATION',(1350,190,2100,700),['PK VehicleDelegationId','FK VehicleId','FK FromEmployeeId','FK ToEmployeeId','Status']),('GATE',(70,900,520,1230),['PK GateId','GateName','Location']),('ACCESS_LOG',(690,850,1160,1280),['PK AccessLogId','FK GateId','EmployeeId?','Timestamp']),('ZONE_TRANSIT',(1350,820,2100,1320),['PK ZoneTransitId','FK EmployeeId','FK SecurityZoneId','FK AccessLogId?','Direction','Timestamp'])],[('EMPLOYEE','VEHICLE',False),('VEHICLE','VEHICLE_DELEGATION',False),('GATE','ACCESS_LOG',False),('ACCESS_LOG','ZONE_TRANSIT',False)]),
'Hình 4.14.':erd('fig_4_14_erd_visitor_hr.png','ERD - Visitor & HR',[('EMPLOYEE',(70,230,520,620),['PK EmployeeId','Name','Department']),('PRE_REGISTRATION',(680,180,1240,670),['PK RegistrationId','FK GuestId','FK HostEmployeeId','ExpectedTimeIn','ExpectedTimeOut']),('VISITOR_DETAIL',(1450,230,2050,620),['PK VisitorDetailId','FK RegistrationId','VisitorName']),('WORK_SCHEDULE',(70,850,600,1280),['PK ScheduleId','FK EmployeeId','FK ShiftId','WorkDate']),('SHIFT',(750,880,1180,1240),['PK ShiftId','Name','StartTime','EndTime']),('ATTENDANCE',(1350,820,1900,1320),['PK AttendanceId','FK EmployeeId','FK ScheduleId?','WorkDate','CheckIn / CheckOut']),('LEAVE_REQUEST',(1950,800,2340,1320),['PK LeaveRequestId','FK EmployeeId','ApproverId?','StartDate / EndDate','Status'])],[('EMPLOYEE','PRE_REGISTRATION',False),('PRE_REGISTRATION','VISITOR_DETAIL',False),('EMPLOYEE','WORK_SCHEDULE',False),('SHIFT','WORK_SCHEDULE',False),('WORK_SCHEDULE','ATTENDANCE',False),('EMPLOYEE','LEAVE_REQUEST',False)]),
'Hình 4.15.':erd('fig_4_15_erd_soc_evidence.png','ERD - SOC & Evidence',[('ALARM',(50,180,480,580),['PK AlarmId','Severity','State','AssignedToUserId?']),('ALARM_COMMENT',(610,180,1080,580),['PK AlarmCommentId','FK AlarmId','UserId?','Comment']),('INCIDENT',(1210,180,1680,580),['PK IncidentId','PrimaryAlarmId?','OwnerUserId?','Status']),('DISPATCH_TASK',(1810,180,2320,580),['PK DispatchTaskId','AlarmId?','IncidentId?','AssignedGuardUserId?']),('EVIDENCE_ITEM',(200,850,760,1330),['PK EvidenceItemId','AlarmId?','IncidentId?','HashSha256','IntegrityStatus']),('REDACTION_REQUEST',(900,850,1470,1330),['PK RedactionRequestId','FK EvidenceItemId','Status']),('EXPORT_REQUEST',(1610,850,2180,1330),['PK EvidenceExportRequestId','EvidenceItemId?','Status','ExportHash?'])],[('ALARM','ALARM_COMMENT',False),('ALARM','INCIDENT',True),('INCIDENT','DISPATCH_TASK',True),('INCIDENT','EVIDENCE_ITEM',True),('EVIDENCE_ITEM','REDACTION_REQUEST',False),('EVIDENCE_ITEM','EXPORT_REQUEST',True)]),
'Hình 4.16.':classdiag()}

doc=Document(SRC)
captions=[p for p in doc.paragraphs if p.text.strip().startswith('Hình ')]
for cap in captions:
 key=next((k for k in assets if cap.text.strip().startswith(k)),None)
 if not key: continue
 prev=cap._p.getprevious()
 while prev is not None and not prev.xpath('.//w:drawing'): prev=prev.getprevious()
 if prev is None: continue
 from docx.text.paragraph import Paragraph
 pp=Paragraph(prev,cap._parent);pp.clear();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.keep_with_next=True;pp.add_run().add_picture(str(assets[key]),width=Inches(6.25))
 if key=='Hình 4.12.':cap.text='Hình 4.12. Mô hình miền dữ liệu V-Shield 2.0';cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
for p in doc.paragraphs:
 if p.text.strip()=='4.3.1 ERD tổng quan':p.text='4.3.1 Mô hình miền dữ liệu';p.style='Heading 3'
 if p.text.strip()=='4.4 Class / Repository Diagram':p.text='4.4 UML Class Diagram - Gate Transit';p.style='Heading 2'
 if p.text.strip().startswith('Hình 4.16.'):
  p.text='Hình 4.16. UML Class Diagram module Gate Transit';p.alignment=WD_ALIGN_PARAGRAPH.CENTER

# SRS mục tiêu/hậu điều kiện cụ thể, không lặp câu template.
objectives={
'UC-02':('Kiểm soát lượt người và phương tiện qua cổng bằng QR, ANPR và chính sách truy cập.','Lượt ra/vào và quyết định Allow/Deny được lưu; barrier chỉ mở khi điều kiện hợp lệ.'),
'UC-03':('Cho phép bảo vệ xử lý ngoại lệ tại cổng khi giao dịch bị từ chối nhưng có căn cứ cho qua, đồng thời lưu đầy đủ lý do và trách nhiệm người thực hiện.','Barrier được mở theo override; lý do, actor, timestamp và Receipt ID được lưu vào audit trail.'),
'UC-04':('Phát cảnh báo cưỡng ép bí mật tới SOC và bảo vệ cơ động mà không tạo tín hiệu gây nguy hiểm cho bảo vệ tại cổng.','Critical Alarm được tạo và phân phối realtime đến SOC/mobile; trạng thái acknowledge được lưu.'),
'UC-05':('Tiếp nhận, phân công, theo dõi và đóng sự cố an ninh theo một vòng đời xử lý có kiểm soát.','Alarm/Incident được cập nhật trạng thái, người phụ trách, timeline và kết quả xử lý.'),
'UC-06':('Tạo lịch hẹn khách có thời gian, host và thông tin cần thiết trước khi khách đến.','PreRegistration được tạo và khách nhận token/link hoặc QR sau khi đáp ứng điều kiện phê duyệt.'),
'UC-07':('Xác thực đăng ký và ghi nhận khách đến tại kiosk hoặc với hỗ trợ của lễ tân.','Thời điểm check-in được lưu và Visitor Pass được phát khi QR, thời gian và phê duyệt hợp lệ.'),
'UC-08':('Cho phép chủ xe giao quyền sử dụng phương tiện cho đồng nghiệp theo một yêu cầu có trạng thái rõ ràng.','VehicleDelegation được lưu ở trạng thái Pending, Approved, Rejected hoặc Revoked cùng thời điểm phản hồi.'),
'UC-09':('Cho phép nhân viên gửi yêu cầu nghỉ phép và quản lý xử lý trực tuyến.','LeaveRequest được lưu với trạng thái và người phê duyệt; kết quả có thể dùng khi tổng hợp công.'),
'UC-10':('Tổng hợp lượt qua vùng và lịch làm việc để tính thời gian làm, đi trễ, về sớm và bất thường.','Attendance được tạo hoặc tính lại; các bất thường thiếu lượt hoặc sai lịch được ghi nhận để rà soát.'),
'UC-11':('Đăng ký bằng chứng số, bảo vệ tính toàn vẹn và cho phép kiểm tra lại SHA-256.','Evidence được đăng ký kèm SHA-256 và trạng thái integrity có thể được xác minh lại.'),
'UC-12':('Che dữ liệu cá nhân và xuất bằng chứng theo quy trình phê duyệt, bảo toàn hash và Chain of Custody.','Bản redacted/export được tạo sau phê duyệt, kèm hash và bản ghi custody tương ứng.')}
for t in doc.tables:
 if len(t.columns)!=2:continue
 vals={r.cells[0].text.strip():r.cells[1] for r in t.rows}
 uid=vals.get('Use Case ID')
 if uid and uid.text.strip() in objectives:
  o,post=objectives[uid.text.strip()];vals['Mục tiêu'].text=o;vals['Hậu điều kiện'].text=post
 # entity detail: prevent bad word breaking
 if t.rows and t.rows[0].cells[0].text.strip()=='Thực thể':
  widths=[2.0,1.25,.75,1.5,.8];t.autofit=False
  grid=t._tbl.tblGrid
  for gc,w in zip(grid.gridCol_lst,widths):gc.set(qn('w:w'),str(int(w*1440)))
  for row in t.rows:
   for i,c in enumerate(row.cells):
    c.width=Inches(widths[i]);
    tcW=c._tc.get_or_add_tcPr().find(qn('w:tcW'))
    if tcW is None:tcW=OxmlElement('w:tcW');c._tc.get_or_add_tcPr().append(tcW)
    tcW.set(qn('w:w'),str(int(widths[i]*1440)));tcW.set(qn('w:type'),'dxa')
    if i==0 and row is not t.rows[0]:
     raw=c.text.replace('\u2060','');display={'VehicleDelegation':'Vehicle Delegation','PreRegistration':'Pre Registration','LeaveRequest':'Leave Request','EvidenceItem':'Evidence Item','RedactionRequest':'Redaction Request','ChainOfCustodyEntry':'Chain of Custody Entry'}.get(raw,raw);c.text=display
    for p in c.paragraphs:
     p.paragraph_format.keep_together=True
     for r in p.runs:
      r.font.size=Pt(7.5)
      if i==0:r.font.name='Arial Narrow';r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'),'Arial Narrow');r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'),'Arial Narrow')

def replace_interaction_table(t,rows):
 while len(t.rows)<len(rows):t.add_row()
 for ri,row in enumerate(t.rows):
  vals=rows[ri] if ri<len(rows) else ['']*4
  for ci,c in enumerate(row.cells):c.text=vals[ci] if ci<len(vals) else ''
for t in doc.tables:
 flat=' | '.join(c.text for r in t.rows for c in r.cells)
 if 'Camera Live' in flat and 'ANPR Result' in flat:
  replace_interaction_table(t,[['TT','Điều khiển','Sự kiện','Mô tả hoạt động'],['1','Camera Live','Initialize','Mở luồng camera tại gate'],['2','ANPR Result','Detect','Nhận biển số và confidence'],['3','QR Validation','Scan','Xác thực trạng thái QR'],['4','Subject Detail','Select','Hiển thị nhân viên/khách và phương tiện'],['5','Decision','Evaluate','Hiển thị Allow/Deny và lý do'],['6','Open Barrier','Click','Mở barrier khi quyết định Allow'],['7','Override','Click','Yêu cầu lý do và ghi audit/receipt'],['8','Duress','Click','Tạo Critical Alarm kín đáo']])
 elif 'Critical alert' in flat and 'Open Incident' in flat:
  replace_interaction_table(t,[['TT','Điều khiển','Sự kiện','Mô tả hoạt động'],['1','Critical badge','Receive','Hiển thị mức độ Critical'],['2','Alarm title','Render','Hiển thị loại Duress Alarm'],['3','Gate / Time / Source','Render','Hiển thị vị trí, timestamp và nguồn'],['4','Incident Detail','Open','Hiển thị nội dung và hướng dẫn'],['5','Acknowledge','Tap','Ghi nhận tiếp nhận và dừng chuông/rung'],['6','Open Incident','Tap','Chuyển sang màn hình xử lý sự cố']])
 elif 'Evidence list' in flat and 'Custody Log' in flat:
  replace_interaction_table(t,[['TT','Điều khiển','Sự kiện','Mô tả hoạt động'],['1','KPI cards','Initialize','Tổng hợp evidence và yêu cầu chờ xử lý'],['2','Evidence table','Load','Tải danh sách theo bộ lọc và quyền'],['3','Detail drawer','Select','Hiển thị metadata, hash và custody'],['4','Verify','Click','Đối chiếu SHA-256'],['5','Redaction','Click','Tạo/duyệt yêu cầu che dữ liệu'],['6','Export','Click','Tạo yêu cầu xuất có phê duyệt'],['7','Custody','Open','Xem lịch sử truy cập/chuyển giao']])

settings=doc.settings._element;upd=settings.find(qn('w:updateFields'))
if upd is None:upd=OxmlElement('w:updateFields');settings.append(upd)
upd.set(qn('w:val'),'true');doc.save(OUT);print(OUT)
