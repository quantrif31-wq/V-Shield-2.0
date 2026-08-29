from pathlib import Path
from copy import deepcopy
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r'C:\Code\V-Shield-2.0')
SRC=Path(r'C:\Users\AnhPhamPC\Downloads\PRO2191 - Dự án TN (UDPM-.NET Core)\BaoCaoTotNghiep.docx')
OUT=SRC.parent/'BaoCaoTotNghiep_VShield_Chapter3_4_Updated.docx'
ASSET=ROOT/'tools'/'report_assets'; ASSET.mkdir(exist_ok=True)
FONT=r'C:\Windows\Fonts\arial.ttf'; FONTB=r'C:\Windows\Fonts\arialbd.ttf'

def font(n,b=False): return ImageFont.truetype(FONTB if b else FONT,n)
def wrap(draw,text,f,maxw):
    words=text.split(); lines=[]; cur=''
    for w in words:
        test=(cur+' '+w).strip()
        if draw.textbbox((0,0),test,font=f)[2] <= maxw: cur=test
        else:
            if cur: lines.append(cur)
            cur=w
    if cur: lines.append(cur)
    return lines
def box(draw,xy,title,items=(),fill='#F4F7FA',outline='#24557A',title_fill=None,fs=25):
    x1,y1,x2,y2=xy; draw.rounded_rectangle(xy,16,fill=fill,outline=outline,width=3)
    if title_fill:
        draw.rounded_rectangle((x1,y1,x2,y1+55),16,fill=title_fill,outline=outline,width=3)
    draw.text((x1+18,y1+13),title,font=font(fs,True),fill='#153047')
    y=y1+68
    for item in items:
        for line in wrap(draw,'• '+item,font(fs-5),x2-x1-34):
            draw.text((x1+18,y),line,font=font(fs-5),fill='#172B3A'); y+=fs
        y+=5
def arrow(draw,a,b,color='#50758F',w=5):
    draw.line((a,b),fill=color,width=w); x,y=b; draw.polygon([(x,y),(x-14,y-8),(x-14,y+8)],fill=color)
def canvas(title,w=1800,h=1100):
    im=Image.new('RGB',(w,h),'white'); d=ImageDraw.Draw(im)
    d.text((60,35),title,font=font(39,True),fill='#14324A'); d.line((60,90,w-60,90),fill='#2E6F95',width=5)
    return im,d
def save(im,name):
    p=ASSET/name; im.save(p,dpi=(220,220),quality=95); return p
def columns(name,title,cols):
    im,d=canvas(title); gap=35; cw=(1680-gap*(len(cols)-1))//len(cols)
    for i,(head,items) in enumerate(cols):
        x=60+i*(cw+gap); box(d,(x,145,x+cw,1010),head,items,title_fill='#DCEAF3',fs=25)
        if i<len(cols)-1: arrow(d,(x+cw+5,580),(x+cw+gap-5,580))
    return save(im,name)
def actors_diagram(name,title,groups):
    im,d=canvas(title); actor_x=75; use_x=430
    actors=list(groups); ah=820//max(1,len(actors))
    for i,(actor,uses) in enumerate(groups.items()):
        cy=165+i*ah+ah//2; d.ellipse((actor_x+45,cy-65,actor_x+85,cy-25),outline='#173B54',width=4)
        d.line((actor_x+65,cy-25,actor_x+65,cy+35),fill='#173B54',width=4); d.line((actor_x+25,cy,actor_x+105,cy),fill='#173B54',width=4); d.line((actor_x+65,cy+35,actor_x+30,cy+85),fill='#173B54',width=4); d.line((actor_x+65,cy+35,actor_x+100,cy+85),fill='#173B54',width=4)
        d.text((actor_x,cy+92),actor,font=font(20,True),fill='#173B54')
        for j,u in enumerate(uses):
            x=use_x+(j%3)*420; y=145+i*ah+(j//3)*80
            d.rounded_rectangle((x,y,x+360,y+58),28,fill='#F3F8FB',outline='#2E6F95',width=3); d.text((x+15,y+15),u,font=font(18),fill='#16364D'); d.line((actor_x+110,cy+10,x,y+29),fill='#95A9B7',width=2)
    return save(im,name)
def wire(name,title,left,right,accent='#2E6F95'):
    im,d=canvas(title); box(d,(60,135,1735,215),'TOPBAR / TRẠNG THÁI HỆ THỐNG',(),fill='#EAF1F6',outline=accent,fs=24)
    box(d,(60,245,390,1010),left[0],left[1:],fill='#F6F8FA',outline=accent,fs=23)
    n=len(right); hh=(745-25*(n-1))//n
    for i,(h,*items) in enumerate(right): box(d,(425,245+i*(hh+25),1735,245+i*(hh+25)+hh),h,items,fill='white',outline=accent,fs=23)
    return save(im,name)
def tree(name,title,roots):
    im,d=canvas(title); y=140
    for root,children in roots:
        box(d,(70,y,430,y+70),root,(),fill='#DCEAF3',fs=24); cy=y
        for j,ch in enumerate(children):
            yy=y+j*65; d.line((430,y+35,520,yy+25),fill='#6C879A',width=3); d.rounded_rectangle((520,yy,1690,yy+50),12,fill='#F7FAFC',outline='#7EA0B7',width=2); d.text((540,yy+12),ch,font=font(20),fill='#18364A')
        y += max(110,len(children)*65+35)
    return save(im,name)

# Diagram assets
overview=columns('fig31.png','Mô hình tổng quan hệ thống V-Shield 2.0',[
 ('Người dùng',['Admin; Quản lý/Nhân sự','Bảo vệ cổng; Bảo vệ cơ động','Nhân viên; Lễ tân; Khách']),
 ('Kênh tương tác',['Vue 3 Web Admin/SOC','Gate Transit Console','Kotlin Android App','Visitor/Kiosk UI']),
 ('Nền tảng V-Shield',['ASP.NET Core 8 Modular Monolith','JWT/RBAC; SignalR realtime','Gate, SOC, Visitor, HR, Evidence']),
 ('Hạ tầng',['SQL Server','Edge AI/ANPR; QR','Camera; Barrier; thiết bị cổng'])])
uc_over=actors_diagram('fig32.png','Use Case tổng quan',{
 'Admin':['Administration','Evidence','Device'], 'Quản lý/Nhân sự':['HR & Attendance','SOC','Approvals'], 'Bảo vệ':['Gate Transit','SOC Alarm','Override'], 'Nhân viên/Khách':['Visitor','Vehicle','Self-service']})
uc_gate=actors_diagram('fig33.png','Use Case Gate Transit & SOC',{
 'Bảo vệ cổng':['Theo dõi transit','Scan QR','Nhận diện ANPR','Open Barrier','Override + lý do','Duress Alarm'], 'Bảo vệ cơ động':['Receive Alarm','Acknowledge','Incident'], 'Security Manager':['Assign Incident','Close Incident','Review Receipt']})
uc_vis=actors_diagram('fig34.png','Use Case Visitor & Employee Self-Service',{
 'Nhân viên':['Pre-register Visitor','Register Vehicle','Delegate Vehicle','Leave Request','View Attendance'], 'Khách':['Self Registration','Visitor QR','Kiosk Check-in'], 'Lễ tân/Quản lý':['Approve Visitor','Issue Pass','Approve Request']})
uc_admin=actors_diagram('fig35.png','Use Case Administration, HR & Evidence',{
 'Admin':['Users/RBAC','Device Topology','Notification Rules','Audit Logs','Full Lockdown'], 'Quản lý':['SOC/Incident','Review Export','Evidence'], 'Nhân sự':['Employee','Shift','Attendance','Leave Approval']})
arch=columns('fig41.png','Kiến trúc công nghệ V-Shield 2.0',[
 ('Presentation',['Vue 3 Web Admin/SOC','Vue 3 Gate Console','Kotlin Android App','Visitor/Kiosk UI']),
 ('Giao tiếp',['HTTPS REST','JWT/RBAC','SignalR realtime']),
 ('ASP.NET Core 8',['Auth; Gate Transit; SOC','Visitor; HR/Attendance','Vehicle; Evidence; Notification']),
 ('Infrastructure',['EF Core + SQL Server','Edge AI / ANPR','Camera; Barrier; QR Device'])])
sitemap=tree('fig42.png','Sitemap V-Shield 2.0',[
 ('WEB ADMIN / SOC',['Dashboard | SOC: Active Alarms, Incident Detail, SOP','Gate Control: Transit, Barrier, History','Visitor: Pre-registration, Pass, Check-in, Overstay, Watchlist','HR & Attendance: Employee, Shift, Attendance, Anomaly, Leave','Evidence: Repository, Verification, Redaction, Export, Custody','Devices: Topology | Administration: Users, Roles, Rules, Audit']),
 ('MOBILE',['Home | My QR | Alerts | Incident | Visitor | Vehicle | Attendance | Leave | Profile']),
 ('KIOSK',['Welcome | Scan QR/Identification | Check-in | Confirmation/Visitor Pass'])])
layouts=[
 wire('fig43.png','Layout 1 - Web Admin / SOC',('LEFT NAV','Dashboard','SOC','Gate Control','Visitor','HR & Attendance','Evidence','Devices','Administration'),[('KPI / FILTER','Severity, trạng thái, site, thời gian'),('MAIN WORKSPACE','Danh sách và chi tiết nghiệp vụ')]),
 wire('fig44.png','Layout 2 - Gate Transit Console',('GATE STATUS','Lane','Barrier','Device health'),[('LIVE CAMERA + ANPR','Video, biển số, độ tin cậy'),('QR + IDENTITY','Trạng thái QR, nhân viên/khách, xe'),('DECISION','Kết quả, Open Gate, Override, Duress')]),
 wire('fig45.png','Layout 3 - Android Mobile App',('APP BAR','User','Online'),[('CRITICAL ALERT','Severity, gate, timestamp'),('INCIDENT','Tóm tắt và hướng dẫn xử lý'),('ACTIONS','Acknowledge, Open Incident, Bottom navigation')],accent='#A33A3A'),
 wire('fig46.png','Layout 4 - Visitor Kiosk',('WELCOME','Ngôn ngữ','Hỗ trợ'),[('QR SCANNER','Vùng quét và hướng dẫn'),('CHECK-IN','Thông tin xác nhận'),('CONFIRMATION','Visitor Pass / liên hệ lễ tân')])]
screens=[
 wire('fig47.png','Screen 01 - SOC Alarm Console',('FILTERS','① KPI Critical/High','② Active Alarm list','③ Severity'),[('INCIDENT DETAIL','④ Camera/evidence preview','Gate, thời gian, nguồn'),('ACTIONS','⑤ Acknowledge','⑥ Assign','⑦ Close')],accent='#A33A3A'),
 wire('fig48.png','Screen 02 - Gate Transit Console',('GATE','① Live Camera','② ANPR','③ QR Status'),[('SUBJECT','④ Employee/Visitor','⑤ Vehicle','⑥ Result'),('ACTIONS','⑦ Open Gate','⑧ Override','⑨ Duress')]),
 wire('fig49.png','Screen 03 - Visitor Pre-registration / Kiosk',('FLOW','① Token/QR','② Visitor identity','③ Host'),[('VISIT','④ Expected time','⑤ Vehicle/plate'),('ACTIONS','⑥ Submit/Check-in','⑦ Visitor Pass')]),
 wire('fig410.png','Screen 04 - Mobile Critical Alert',('ALERT','① Critical severity','② Gate/location'),[('INCIDENT','③ Nội dung sự cố','④ Timestamp'),('ACTIONS','⑤ Acknowledge','⑥ Open Incident')],accent='#A33A3A'),
 wire('fig411.png','Screen 05 - Evidence Repository',('FILTERS','① Type','② Incident','③ Custody status'),[('EVIDENCE LIST','④ Hash status','⑤ Verification'),('ACTIONS','⑥ Redaction','⑦ Export','⑧ Custody log')])]
erd=columns('fig412.png','ERD tổng quan theo domain',[
 ('Identity',['AppUser','Employee','Role/permission']),('Access/Visitor',['Gate - AccessLog - Vehicle','ZoneTransit - Attendance','PreRegistration - VisitorDetail']),('SOC',['Alarm - Incident','AlarmComment','DispatchTask']),('Evidence',['EvidenceItem','RedactionRequest','EvidenceExportRequest','ChainOfCustodyEntry'])])
erd1=columns('fig413.png','ERD chi tiết - Access & Vehicle',[
 ('Gate',['PK GateId','1 - N AccessLog']),('AccessLog / ZoneTransit',['PK AccessLogId / ZoneTransitId','FK GateId, EmployeeId','Source, Direction, Timestamp']),('Vehicle',['PK VehicleId','FK EmployeeId','UQ LicensePlate']),('VehicleDelegation',['PK VehicleDelegationId','FK VehicleId, From/ToEmployeeId','Status'])])
erd2=columns('fig414.png','ERD chi tiết - Visitor & HR',[
 ('PreRegistration',['PK RegistrationId','FK GuestId, HostEmployeeId']),('VisitorDetail',['PK VisitorDetailId','FK RegistrationId']),('Employee / Shift',['PK EmployeeId / ShiftId']),('Attendance / LeaveRequest',['PK AttendanceId / LeaveRequestId','FK EmployeeId','WorkDate / StartDate-EndDate'])])
erd3=columns('fig415.png','ERD chi tiết - SOC & Evidence',[
 ('Alarm',['PK AlarmId','Severity, State','N - 0..1 Incident']),('Incident',['PK IncidentId','PrimaryAlarmId','OwnerUserId']),('EvidenceItem',['PK EvidenceItemId','FK AlarmId, IncidentId','HashSha256']),('Compliance',['RedactionRequest','EvidenceExportRequest','ChainOfCustodyEntry','FK EvidenceItemId'])])
classdiag=columns('fig416.png','Class / Repository Diagram - Gate Transit',[
 ('GateTransitController',['POST scan / scan-guest','GET gates / logs']),('ZoneTransitService',['Xác thực policy','Ghi transit/attendance','Sinh quyết định']),('ApplicationDbContext',['DbSet<ZoneTransit>','DbSet<Gate/Vehicle/AccessLog>']),('SQL Server + Devices',['Persistence','Camera/ANPR/QR','Barrier'])])

doc=Document(SRC)
body=doc.element.body
p3=next(p for p in doc.paragraphs if p.text.strip()=='PHẦN 3: PHÂN TÍCH - ANALYSIS')
p5=next(p for p in doc.paragraphs if p.text.strip()=='PHẦN 5: THỰC HIỆN – IMPLEMENT')
cur=p3._p
while cur is not p5._p:
    nxt=cur.getnext(); body.remove(cur); cur=nxt

def move(el): body.insert(body.index(p5._p),el)
def para(text='',style=None,bold=False,center=False,keep=False):
    p=doc.add_paragraph(style=style); r=p.add_run(text); r.bold=bold
    if center:p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if keep:p.paragraph_format.keep_with_next=True
    move(p._p); return p
def heading(text,level): return para(text,f'Heading {level}',keep=True)
def pagebreak():
    p=doc.add_paragraph(); p.add_run().add_break(); p.runs[0]._element.getparent().getparent(); p._p.get_or_add_pPr().append(OxmlElement('w:pageBreakBefore')); move(p._p)
def figure(path,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True; p.add_run().add_picture(str(path),width=Inches(6.35)); move(p._p)
    c=doc.add_paragraph(style='Caption' if 'Caption' in [s.name for s in doc.styles] else None); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.keep_with_next=True; c.add_run(caption).bold=True; move(c._p)
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def table(rows,widths=None,header=True):
    t=doc.add_table(rows=0,cols=len(rows[0])); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; cells[ci].text=str(val)
            if widths: cells[ci].width=Inches(widths[ci])
            if ri==0 and header:
                shade(cells[ci],'D9EAF4')
                for run in cells[ci].paragraphs[0].runs: run.bold=True
        t.rows[-1]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    if header:
        trPr=t.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); trPr.append(hdr)
    move(t._tbl); return t
def prose(text): para(text,'Normal')

heading('PHẦN 3: PHÂN TÍCH - ANALYSIS',1)
heading('3.1 Mô hình tổng quan hệ thống',2); figure(overview,'Hình 3.1. Mô hình tổng quan hệ thống V-Shield 2.0')
prose('V-Shield 2.0 là nền tảng kiểm soát vào ra và vận hành an ninh doanh nghiệp. Người dùng truy cập qua Web Admin/SOC, bàn điều khiển tại cổng, ứng dụng Android và kiosk khách. Các kênh gửi yêu cầu HTTPS REST với JWT/RBAC; sự kiện khẩn cấp và trạng thái vận hành được đẩy thời gian thực bằng SignalR. Backend ASP.NET Core 8 được tổ chức theo Modular Monolith, dùng Entity Framework Core và SQL Server, đồng thời tích hợp Edge AI/ANPR, camera, thiết bị QR và barrier.')
heading('3.2 Sơ đồ Use Case',2)
for h,img,cap,desc in [
 ('3.2.1 Tổng quan',uc_over,'Hình 3.2. Use Case tổng quan','Sơ đồ tổng quan thể hiện actor nào tham gia phân hệ nào, không liệt kê toàn bộ thao tác nhỏ. Admin phụ trách quản trị, thiết bị và bằng chứng; Quản lý/Nhân sự tham gia phê duyệt, HR và SOC; lực lượng bảo vệ vận hành cổng, tiếp nhận cảnh báo và xử lý ngoại lệ; nhân viên, lễ tân và khách sử dụng các luồng tự phục vụ về khách, phương tiện, nghỉ phép và chấm công.'),
 ('3.2.2 Gate Transit & SOC',uc_gate,'Hình 3.3. Use Case Gate Transit & SOC','Phạm vi gồm kiểm soát giao dịch qua cổng và vòng đời cảnh báo SOC. Nhân viên bảo vệ tại cổng theo dõi camera, ANPR, QR và quyết định mở barrier. Override chỉ dùng khi luồng bình thường gặp ngoại lệ, bắt buộc ghi lý do và tạo dấu vết trách nhiệm. Duress tạo cảnh báo mức Critical; SOC hoặc bảo vệ cơ động tiếp nhận, acknowledge, được phân công và đóng incident khi đã xử lý.'),
 ('3.2.3 Visitor & Employee Self-Service',uc_vis,'Hình 3.4. Use Case Visitor & Employee Self-Service','Sơ đồ mô tả luồng đăng ký khách trước, khách tự khai báo, nhận QR và check-in tại kiosk; lễ tân hoặc quản lý xử lý phê duyệt và phát hành Visitor Pass. Nhân viên đồng thời quản lý phương tiện, ủy quyền xe, tạo yêu cầu nghỉ phép và xem chấm công cá nhân. Ngoại lệ đáng chú ý là QR hết hạn/không hợp lệ, khách thiếu phê duyệt hoặc ủy quyền xe bị từ chối.'),
 ('3.2.4 Administration, HR & Evidence',uc_admin,'Hình 3.5. Use Case Administration, HR & Evidence','Phạm vi bao gồm quản trị người dùng/RBAC, topology thiết bị, notification rules, audit logs; nghiệp vụ nhân sự, ca làm việc, chấm công, bất thường và nghỉ phép; cùng kho bằng chứng số. Bằng chứng được đăng ký với SHA-256, kiểm tra tính toàn vẹn, xử lý redaction và xuất sau phê duyệt, đồng thời ghi Chain of Custody. Full Lockdown là thao tác đặc quyền và phải được audit.')]:
    heading(h,3); figure(img,cap); prose(desc)

heading('3.3 Đặc tả yêu cầu hệ thống (SRS)',2)
specs=[
('UC-01','Đăng nhập hệ thống','Người dùng','Xác thực và cấp phiên JWT theo vai trò','Người dùng mở màn hình đăng nhập','Tài khoản hoạt động','JWT/refresh token được cấp; quyền được nạp','Nhập thông tin; hệ thống xác thực; kiểm tra trạng thái/MFA; phát JWT; chuyển tới màn hình theo vai trò','MFA hoặc refresh token khi phiên hết hạn','Sai thông tin, tài khoản khóa, thiếu quyền','AppUser, UserRefreshToken','Không lộ lý do xác thực chi tiết; audit thất bại.'),
('UC-02','Kiểm soát người/phương tiện vào ra','Bảo vệ cổng','QR device, ANPR camera, barrier','Có lượt tiếp cận cổng','Gate và thiết bị online; subject có dữ liệu','Nhận QR/biển số; gọi scan; đối chiếu người-xe-quyền; trả quyết định; mở barrier nếu hợp lệ; ghi AccessLog/ZoneTransit và Receipt ID','Tra cứu thủ công subject hoặc biển số','QR hết hạn, biển số không khớp, anti-passback/policy từ chối','Gate, Vehicle, AccessLog, ZoneTransit','Chỉ mở barrier khi quyết định Allow hoặc override hợp lệ.'),
('UC-03','Override mở cổng','Bảo vệ cổng','Quản lý/SOC','Giao dịch bình thường bị từ chối nhưng cần xử lý ngoại lệ','Người dùng có quyền override; gate xác định','Chọn Override; nhập lý do; xác nhận trách nhiệm; backend kiểm tra quyền; mở barrier; tạo receipt/audit','Quản lý phê duyệt theo policy','Thiếu lý do, không đủ quyền, barrier offline','AccessLog, ExceptionReason, SystemAuditLog','Lý do bắt buộc; không sửa/xóa dấu vết override.'),
('UC-04','Kích hoạt Duress Alarm','Bảo vệ cổng','SOC, bảo vệ cơ động','Người vận hành nhấn Duress','Đã xác định gate và phiên người dùng','Gửi duress; backend ghi sự kiện; tạo Alarm Critical; SignalR đẩy SOC/mobile; app rung/chuông; người nhận acknowledge','Gửi kèm vị trí hoặc source device','Mất mạng: ghi cục bộ/hiển thị thất bại để gọi kênh dự phòng','DuressEvent, Alarm, Notification','Duress luôn Critical; không hiển thị phản hồi gây nguy hiểm tại cổng.'),
('UC-05','Tiếp nhận và xử lý cảnh báo SOC','SOC / Security Manager','Bảo vệ cơ động','Alarm mới được tạo','SOC online; người dùng có quyền','Nhận realtime; mở chi tiết; acknowledge; tạo/gắn incident; assign; comment; thực hiện SOP; close với outcome','Chuyển người xử lý hoặc nâng severity','Alarm trùng, mất kết nối, thiếu outcome','Alarm, Incident, AlarmComment, DispatchTask','Mọi thay đổi trạng thái có timestamp và actor.'),
('UC-06','Đăng ký khách trước','Nhân viên (Host)','Khách, lễ tân','Host tạo lịch hẹn khách','Host đăng nhập và còn hiệu lực','Nhập khách, thời gian, số lượng/xe; tạo link/token; khách khai báo; hệ thống lưu; phát QR/pass sau điều kiện phê duyệt','Lễ tân cập nhật trạng thái','Token hết hạn, thời gian không hợp lệ, watchlist/risk','PreRegistration, GuestProfile, VisitorDetail','ExpectedTimeOut sau ExpectedTimeIn; token không tái sử dụng trái phép.'),
('UC-07','Visitor/Kiosk Check-in','Khách','Lễ tân, QR device','Khách quét QR tại kiosk','PreRegistration hợp lệ và trong cửa sổ thời gian','Quét QR; verify; hiển thị thông tin tối thiểu; xác nhận check-in; phát Visitor Pass; ghi access','Nhờ lễ tân hỗ trợ nhận dạng thủ công','QR không hợp lệ/hết hạn, chưa phê duyệt','PreRegistration, VisitorDetail, AccessLog','Chỉ hiển thị dữ liệu cần thiết; log thời điểm check-in.'),
('UC-08','Ủy quyền phương tiện','Nhân viên sở hữu xe','Nhân viên nhận ủy quyền','Chủ xe tạo yêu cầu','Xe thuộc người gửi; hai nhân viên hoạt động','Chọn xe/người nhận; nhập lý do; gửi; người nhận approve; hệ thống áp dụng quyền tạm thời','Reject hoặc revoke','Trùng ủy quyền, xe không thuộc người gửi','Vehicle, VehicleDelegation, Employee','Trạng thái Pending/Approved/Rejected/Revoked; lưu thời điểm phản hồi.'),
('UC-09','Tạo yêu cầu nghỉ phép','Nhân viên','Quản lý/Nhân sự','Nhân viên gửi đơn','Đăng nhập; khoảng ngày hợp lệ','Chọn loại nghỉ; nhập ngày/lý do; gửi Pending; approver xem và approve/reject; đồng bộ tính công','Nhân viên cancel trước khi duyệt','Ngày sai, thiếu lý do, xung đột policy','LeaveRequest, Employee, AppUser','Lý do bắt buộc; reject cần lý do; audit approver.'),
('UC-10','Tính chấm công','Nhân sự / Hệ thống','Nhân viên','Chạy derive/recalculate theo ngày','Có Shift/WorkSchedule và ZoneTransit/AccessLog','Tập hợp transit; xác định check-in/out; tính trễ, về sớm, giờ làm/OT; lưu Attendance; phát hiện anomaly','Điều chỉnh có ghi chú hoặc batch','Thiếu transit, lịch ca chồng lấn','Attendance, Shift, WorkSchedule, ZoneTransit, AttendanceAnomaly','Không tự đoán dữ liệu thiếu; điều chỉnh phải có audit.'),
('UC-11','Quản lý và Verify Evidence','Admin/SOC','Hệ thống lưu trữ','Đăng ký hoặc mở bằng chứng','Nguồn tồn tại; người dùng có quyền','Đăng ký metadata; tính/lưu SHA-256; liên kết Alarm/Incident; verify hash; hiển thị trạng thái; ghi custody','Đưa vào collection/legal hold','Hash mismatch, file mất, storage lỗi','EvidenceItem, EvidenceCollection, ChainOfCustodyEntry','Evidence immutable; SHA-256 mismatch phải cảnh báo và audit.'),
('UC-12','Redaction / Export Evidence','Admin/Quản lý','Người phê duyệt','Có yêu cầu chia sẻ bằng chứng','Evidence tồn tại và không bị cấm xuất','Tạo redaction; phê duyệt; thực hiện/verify bản che; tạo export request; phê duyệt; sinh file/hash/watermark; ghi custody','Export collection hoặc từ chối','Chưa redaction, thiếu approval, legal hold/policy cấm','RedactionRequest, EvidenceExportRequest, ChainOfCustodyEntry','Tách vai trò request/approve khi policy yêu cầu; export có hash.')]
for idx,s in enumerate(specs):
    if idx in (0,5,8): heading({0:'3.3.1 Gate/SOC Use Cases',5:'3.3.2 Visitor/Employee Use Cases',8:'3.3.3 Administration/HR/Evidence Use Cases'}[idx],3)
    para(f'{s[0]} - {s[1]}',bold=True,keep=True)
    labels=['Use Case ID','Tên Use Case','Actor chính','Actor phụ','Mục tiêu','Trigger','Tiền điều kiện','Hậu điều kiện','Luồng chính','Luồng thay thế','Ngoại lệ','Dữ liệu liên quan','Business Rule']
    if len(s)==12:
        actor2='Hệ thống xác thực'; objective=s[3]; trigger=s[4]; pre=s[5]; post=s[6]; main=s[7]; alt=s[8]; exc=s[9]; data=s[10]; rule=s[11]
    else:
        actor2=s[3]; objective='Thực hiện '+s[1].lower()+' đúng quyền và có dấu vết kiểm toán.'; trigger=s[4]; pre=s[5]; post='Dữ liệu và trạng thái nghiệp vụ được cập nhật nhất quán.'; main=s[6]; alt=s[7]; exc=s[8]; data=s[9]; rule=s[10]
    vals=[s[0],s[1],s[2],actor2,objective,trigger,pre,post,'1. '+main.replace(';',';\n2. ',1) if ';' in main else '1. '+main,alt,exc,data,rule]
    table([['Thuộc tính','Nội dung']]+list(zip(labels,vals)),[1.55,4.75])

heading('PHẦN 4: THIẾT KẾ - DESIGN',1)
heading('4.1 Mô hình công nghệ',2); figure(arch,'Hình 4.1. Kiến trúc công nghệ V-Shield 2.0')
prose('Lớp trình bày gồm Vue 3 cho Web Admin/SOC và Gate Transit Console, Kotlin Android cho bảo vệ cơ động/nhân viên, cùng Visitor/Kiosk UI. Các client gọi API HTTPS REST, dùng JWT/RBAC để xác thực và phân quyền. SignalR đảm nhiệm cập nhật trạng thái, cảnh báo và điều phối sự cố theo thời gian thực.')
prose('Backend ASP.NET Core 8 được triển khai theo Modular Monolith với các module nghiệp vụ Auth, Gate Transit, SOC, Visitor, Attendance/HR, Vehicle, Evidence và Notification. Entity Framework Core truy cập SQL Server. Tích hợp thiết bị đi qua các adapter/service cho Edge AI/ANPR, camera, QR và barrier; kiến trúc này không giả định microservices.')
heading('4.2 Thiết kế giao diện',2); heading('4.2.1 Sitemap',3); figure(sitemap,'Hình 4.2. Sitemap V-Shield 2.0'); prose('Sitemap phân tách ba bối cảnh sử dụng: Web Admin/SOC cho vận hành tập trung, mobile cho cảnh báo và tự phục vụ cá nhân, kiosk cho check-in khách. Menu được giới hạn theo chức năng đã thể hiện trong Chương 2, API/source và Chương 5; quyền hiển thị thực tế được lọc bằng RBAC.')
heading('4.2.2 Layout',3)
for i,p in enumerate(layouts,1): figure(p,f'Hình 4.{i+2}. Wireframe layout {i}')
prose('Bốn layout ưu tiên tác vụ khác nhau: Web Admin/SOC tối ưu điều hướng và không gian giám sát; Gate Console tập trung quyết định trong vài giây; mobile đặt cảnh báo quan trọng và hành động chính trong vùng dễ chạm; kiosk dùng luồng tuyến tính, chữ lớn và có đường dẫn hỗ trợ thủ công.')
heading('4.2.3 Giao diện chức năng',3)
callouts=[
 [('1','KPI cảnh báo','Initialize','Tổng hợp Critical/High theo bộ lọc'),('2','Active Alarm list','Receive realtime','Nhận Alarm mới qua SignalR'),('3','Incident Detail','Select','Hiển thị nguồn, gate, timeline/evidence'),('4','Acknowledge','Click','Xác nhận đã tiếp nhận'),('5','Assign','Click','Giao incident cho người xử lý'),('6','Close','Click','Đóng với kết quả xử lý')],
 [('1','Camera Live','Initialize','Mở luồng camera tại gate'),('2','ANPR Result','Detect','Nhận biển số và confidence'),('3','QR Validation','Scan','Gọi xác thực QR'),('4','Open Gate','Click','Mở barrier khi Allow'),('5','Override','Click','Yêu cầu lý do và ghi audit'),('6','Duress','Click','Tạo cảnh báo Critical kín đáo')],
 [('1','Token/QR','Scan','Nhận token đăng ký'),('2','Visitor identity','Input','Khai báo/đối chiếu khách'),('3','Host & schedule','Validate','Kiểm tra host và cửa sổ thời gian'),('4','Check-in','Submit','Ghi nhận khách đến'),('5','Visitor Pass','Generate','Hiển thị/phát QR tạm thời')],
 [('1','Critical alert','Receive','Hiển thị severity, gate, timestamp'),('2','Incident','Open','Mở chi tiết và hướng dẫn'),('3','Acknowledge','Tap','Dừng chuông/rung sau khi ghi nhận'),('4','Open Incident','Tap','Chuyển vào màn hình xử lý')],
 [('1','Evidence list','Initialize','Tải danh sách theo quyền'),('2','Hash status','Verify','Đối chiếu SHA-256'),('3','Redaction','Click','Tạo/duyệt yêu cầu che dữ liệu'),('4','Export','Click','Tạo yêu cầu xuất có phê duyệt'),('5','Custody Log','Open','Xem lịch sử truy cập/chuyển giao')]]
for i,(p,rows) in enumerate(zip(screens,callouts),1):
    figure(p,f'Hình 4.{i+6}. Giao diện chức năng {i}')
    table([['TT','Điều khiển','Sự kiện','Mô tả hoạt động']]+rows,[.45,1.45,1.2,3.2])

heading('4.3 Thiết kế dữ liệu',2); heading('4.3.1 ERD tổng quan',3); figure(erd,'Hình 4.12. ERD tổng quan V-Shield 2.0'); prose('ERD tổng quan chỉ giữ các thực thể cốt lõi có căn cứ trong ApplicationDbContext và model source. Quan hệ nghiệp vụ chính đi từ danh tính nhân viên đến lượt qua cổng, phương tiện/ủy quyền, đăng ký khách, cảnh báo-sự cố và bằng chứng. Các thực thể kỹ thuật mở rộng được lược bỏ để bảo đảm khả năng đọc trên khổ A4.')
heading('4.3.2 ERD chi tiết',3)
for p,c in [(erd1,'Hình 4.13. ERD chi tiết Access & Vehicle'),(erd2,'Hình 4.14. ERD chi tiết Visitor & HR'),(erd3,'Hình 4.15. ERD chi tiết SOC & Evidence')]: figure(p,c)
heading('4.3.3 Chi tiết thực thể',3)
dict_rows=[
('ZoneTransit','ZoneTransitId','int','Khóa lượt chuyển vùng','PK'),('ZoneTransit','Direction','string(10)','Chiều IN/OUT','Required, MaxLength(10)'),('VehicleDelegation','Status','string(20)','Trạng thái ủy quyền','Pending/Approved/Rejected/Revoked'),('PreRegistration','ExpectedTimeIn/Out','datetime','Cửa sổ dự kiến','Out sau In'),('Attendance','TotalWorkingHours','decimal(8,2)','Tổng giờ làm','Tính từ transit/schedule'),('LeaveRequest','Reason','string(2000)','Lý do nghỉ','Required'),('Alarm','Severity/State','string(40)','Mức độ/trạng thái cảnh báo','Theo workflow SOC'),('Incident','PrimaryAlarmId','long?','Cảnh báo chính','FK logic, nullable'),('EvidenceItem','HashSha256','string(128)','Giá trị băm SHA-256','MaxLength(128)'),('RedactionRequest','EvidenceItemId','long','Bằng chứng cần che','FK'),('ChainOfCustodyEntry','Action','string(80)','Hành động custody','MaxLength(80)')]
table([['Thực thể','Thuộc tính','Kiểu dữ liệu','Mô tả','Ràng buộc']]+dict_rows,[1.2,1.5,1.0,1.7,.9])
prose('Kiểu dữ liệu và ràng buộc trong bảng được lấy từ model C# và Data Annotation/Column(TypeName) hiện có. Các quan hệ không có ForeignKey hoặc Fluent API được chứng minh rõ chỉ được mô tả ở mức liên kết nghiệp vụ, không khẳng định ràng buộc SQL.')
heading('4.4 Class / Repository Diagram',2); figure(classdiag,'Hình 4.16. Class / Repository Diagram module Gate Transit')
prose('GateTransitController cung cấp endpoint scan, scan-guest và các truy vấn liên quan đến gate/subject. ZoneTransitService thực hiện kiểm tra nghiệp vụ và ghi nhận lượt di chuyển. ApplicationDbContext đóng vai trò lớp truy cập dữ liệu bằng Entity Framework Core với các DbSet thực tế; dữ liệu được lưu trong SQL Server. Tích hợp ANPR, QR và barrier nằm ở lớp service/adapter thiết bị, thay cho mô hình DAO/JPA của template Java/Spring.')

# Loại bỏ dữ liệu kiểm thử thương mại điện tử còn sót trong template Chương 6,
# nhưng giữ nguyên cấu trúc/chủ đề của chương kiểm thử.
for p in doc.paragraphs:
    if 'thay cho mô hình DAO/JPA của template Java/Spring' in p.text:
        for r in p.runs:
            if 'thay cho mô hình DAO/JPA của template Java/Spring' in r.text:
                r.text=r.text.replace('thay cho mô hình DAO/JPA của template Java/Spring','phù hợp với kiến trúc ASP.NET Core hiện tại')
for t in doc.tables:
    all_text=' '.join(c.text for row in t.rows for c in row.cells)
    if 'ProductDate' in all_text or 'OrderDetail' in all_text:
        replacement=[
          ['TT','CHỨC NĂNG','DỮ LIỆU MẪU','KẾT QUẢ MONG ĐỢI','TÌNH TRẠNG'],
          ['1','Gate Transit','','',''],
          ['1.1','QR hợp lệ','QR động còn hạn','Quyết định Allow; ghi AccessLog/Receipt ID',''],
          ['1.2','ANPR không khớp','Biển số khác hồ sơ','Từ chối và hiển thị lý do',''],
          ['1.3','Override','Lý do ngoại lệ hợp lệ','Mở barrier; ghi audit/receipt',''],
          ['1.4','Duress','Gate đang hoạt động','Tạo Alarm Critical và đẩy SignalR',''],
          ['2','Đăng nhập','','',''],
          ['2.1','Để trống','','Yêu cầu nhập','Đã fix'],
          ['2.2','Sai thông tin','Sai username/mật khẩu','Từ chối, không lộ chi tiết nhạy cảm',''],
          ['2.3','Đăng nhập đúng','Tài khoản hoạt động','Cấp JWT và chuyển theo vai trò',''],
          ['2.4','Phân quyền','Role BaoVe Gate','Chỉ hiển thị chức năng được cấp',''],
          ['3','SOC Alarm','','',''],
          ['3.1','Acknowledge','Alarm trạng thái New','Chuyển trạng thái và lưu timestamp',''],
          ['3.2','Close','Incident có outcome','Đóng sự cố và ghi audit','']]
        for ri,row in enumerate(t.rows):
            vals=replacement[ri] if ri<len(replacement) else ['']*len(row.cells)
            for ci,c in enumerate(row.cells): c.text=vals[ci] if ci<len(vals) else ''

# ask Word to refresh TOC/page fields on open
settings=doc.settings._element; upd=settings.find(qn('w:updateFields'))
if upd is None: upd=OxmlElement('w:updateFields'); settings.append(upd)
upd.set(qn('w:val'),'true')
doc.save(OUT)
print(OUT)
